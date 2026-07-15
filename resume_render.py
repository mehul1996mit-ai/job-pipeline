"""Generate a resume as DOCX + PDF from a structured JSON file.

Usage: python make_resume.py <resume.json> <output_basename>
Produces <output_basename>.docx and <output_basename>.pdf
"""
import json
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

ACCENT = "1F3864"  # dark navy


def esc(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_docx(r, path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def para(text="", bold=False, size=10, align=None, space_after=2, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        return p

    def heading(text):
        p = para(text.upper(), bold=True, size=11, space_after=3, color=ACCENT)
        p.paragraph_format.space_before = Pt(8)
        return p

    para(r["name"], bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=ACCENT)
    para(r["contact_line"], size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    heading("Summary")
    para(r["summary"], space_after=4)

    heading("Skills")
    for sk in r["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(sk["label"] + ": ")
        run.bold = True
        p.add_run(sk["items"])

    heading("Professional Experience")
    for exp in r["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(exp["company"])
        run.bold = True
        run.font.size = Pt(11)
        tab = p.add_run("\t" + exp["dates"])
        tab.bold = True
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), 2)  # right-aligned tab
        for role in exp["roles"]:
            rp = doc.add_paragraph()
            rp.paragraph_format.space_after = Pt(2)
            rr = rp.add_run(role["title"])
            rr.bold = True
            rr.italic = True
            if role.get("dates"):
                rd = rp.add_run("\t" + role["dates"])
                rd.italic = True
                rp.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), 2)
            for b in role["bullets"]:
                bp = doc.add_paragraph(b, style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
            for proj in role.get("projects", []):
                pp = doc.add_paragraph()
                pp.paragraph_format.space_after = Pt(2)
                pr = pp.add_run(proj["name"] + ": ")
                pr.bold = True
                pp.add_run(proj["desc"])

    if r.get("achievements"):
        heading("Achievements")
        for a in r["achievements"]:
            ap = doc.add_paragraph(a, style="List Bullet")
            ap.paragraph_format.space_after = Pt(2)

    heading("Education")
    for ed in r["education"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ed["school"])
        run.bold = True
        p.add_run(" -- " + ed["degree"])
        d = p.add_run("\t" + ed["dates"])
        d.bold = True
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), 2)

    doc.save(path)


def build_pdf(r, path):
    doc = SimpleDocTemplate(path, pagesize=A4,
                            topMargin=12 * mm, bottomMargin=12 * mm,
                            leftMargin=15 * mm, rightMargin=15 * mm)
    ss = getSampleStyleSheet()
    name_st = ParagraphStyle("nm", parent=ss["Title"], fontSize=17, spaceAfter=2,
                             textColor="#" + ACCENT, alignment=TA_CENTER)
    contact_st = ParagraphStyle("ct", parent=ss["Normal"], fontSize=8.5,
                                alignment=TA_CENTER, spaceAfter=8)
    head_st = ParagraphStyle("hd", parent=ss["Heading2"], fontSize=10.5,
                             textColor="#" + ACCENT, spaceBefore=8, spaceAfter=2)
    body_st = ParagraphStyle("bd", parent=ss["Normal"], fontSize=9, leading=11.5, spaceAfter=2)
    bullet_st = ParagraphStyle("bl", parent=body_st, leftIndent=12, bulletIndent=4)

    story = [Paragraph(esc(r["name"]), name_st),
             Paragraph(esc(r["contact_line"]), contact_st)]

    def sect(title):
        story.append(Paragraph(title.upper(), head_st))
        story.append(HRFlowable(width="100%", thickness=0.7, color="#" + ACCENT, spaceAfter=4))

    sect("Summary")
    story.append(Paragraph(esc(r["summary"]), body_st))

    sect("Skills")
    for sk in r["skills"]:
        story.append(Paragraph("<b>%s:</b> %s" % (esc(sk["label"]), esc(sk["items"])), body_st))

    sect("Professional Experience")
    for exp in r["experience"]:
        story.append(Paragraph("<b>%s</b> &nbsp;&nbsp; <i>%s</i>"
                               % (esc(exp["company"]), esc(exp["dates"])), body_st))
        for role in exp["roles"]:
            t = "<b><i>%s</i></b>" % esc(role["title"])
            if role.get("dates"):
                t += " &nbsp;&nbsp; <i>%s</i>" % esc(role["dates"])
            story.append(Paragraph(t, body_st))
            for b in role["bullets"]:
                story.append(Paragraph(esc(b), bullet_st, bulletText="•"))
            for proj in role.get("projects", []):
                story.append(Paragraph("<b>%s:</b> %s" % (esc(proj["name"]), esc(proj["desc"])), body_st))
        story.append(Spacer(1, 3))

    if r.get("achievements"):
        sect("Achievements")
        for a in r["achievements"]:
            story.append(Paragraph(esc(a), bullet_st, bulletText="•"))

    sect("Education")
    for ed in r["education"]:
        story.append(Paragraph("<b>%s</b> -- %s &nbsp;&nbsp; <i>%s</i>"
                               % (esc(ed["school"]), esc(ed["degree"]), esc(ed["dates"])), body_st))

    doc.build(story)


def main():
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        r = json.load(f)
    base = sys.argv[2]
    build_docx(r, base + ".docx")
    build_pdf(r, base + ".pdf")
    print("Wrote %s.docx and %s.pdf" % (base, base))


if __name__ == "__main__":
    main()
